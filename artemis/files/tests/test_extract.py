"""Extraction tests.

Deliberately built around the shapes that actually break, not the happy path:
a .tsv whose fields contain commas, single-column files (which defeat
`csv.Sniffer` entirely), Slack's habit of labelling uploads `text/plain` or
`application/octet-stream`, and formats we must refuse BY NAME, not generically.
"""

from __future__ import annotations

import io

import pytest

from artemis.files.extract import extract, supported_extensions
from artemis.files.extract.base import (
    MAX_TABULAR_SAMPLE_ROWS,
    FileParseError,
    FileTooLargeError,
    UnsupportedFileTypeError,
)

# A TSV whose values contain commas, shaped like the demand-gen upload that
# started this. Sniffing happens to get this one right; routing on the extension
# means we do not depend on that holding for the next file.
TSV_WITH_COMMAS = (
    b"District\tState\tEnrollment\tNotes\n"
    b"Austin ISD\tTX\t73000\tPilot, expansion pending\n"
    b"San Diego USD\tCA\t95000\tRenewal, Q3\n"
    b"Denver Public Schools\tCO\t88000\tNew, no contact yet\n"
)


def test_tsv_with_commas_in_values_keeps_its_columns() -> None:
    """Commas inside tab-separated values must not split a column."""
    result = extract(TSV_WITH_COMMAS, filename="demand_gen.tsv", mimetype="text/plain")

    assert result.kind == "tabular"
    table = result.tables[0]
    assert table.columns == ["District", "State", "Enrollment", "Notes"]
    assert table.total_rows == 3
    assert table.sample_rows[0] == ["Austin ISD", "TX", "73000", "Pilot, expansion pending"]


def test_slack_octet_stream_mimetype_does_not_defeat_routing() -> None:
    """Slack labels plenty of uploads application/octet-stream; the name still wins."""
    result = extract(TSV_WITH_COMMAS, filename="export.tsv", mimetype="application/octet-stream")
    assert result.kind == "tabular"
    assert result.tables[0].columns[0] == "District"


def test_single_column_file_survives_a_sniffer_that_cannot_guess() -> None:
    """`csv.Sniffer` raises on a one-column file -- there is no delimiter to find.

    Verified 2026-08-25: sniffing "District\nAustin ISD\n" raises csv.Error.
    Routing on the extension never asks the question, and the comma fallback
    makes the answer irrelevant even for an unknown name.
    """
    result = extract(b"District\nAustin ISD\nDallas ISD\n", filename="districts.csv")
    assert result.tables[0].columns == ["District"]
    assert result.tables[0].total_rows == 2


def test_csv_is_parsed_as_csv() -> None:
    payload = b"name,role\nJosh,Demand Gen\nJon,Marketing Ops"
    result = extract(payload, filename="team.csv")
    assert result.tables[0].columns == ["name", "role"]
    assert result.tables[0].total_rows == 2


def test_large_table_reports_true_row_count_and_says_it_sampled() -> None:
    """An agent must never mistake the sample for the file."""
    rows = b"".join(f"row{i},{i}\n".encode() for i in range(500))
    result = extract(b"name,value\n" + rows, filename="big.csv")

    table = result.tables[0]
    assert table.total_rows == 500
    assert len(table.sample_rows) == MAX_TABULAR_SAMPLE_ROWS
    assert table.truncated
    assert "500" in result.text
    assert "must not be computed from this sample" in result.text


def test_empty_table_fails_loudly_rather_than_returning_nothing() -> None:
    with pytest.raises(FileParseError) as excinfo:
        extract(b"\n\n\n", filename="empty.csv")
    assert "no rows" in excinfo.value.reason


def test_json_is_prettified_and_invalid_json_still_readable() -> None:
    ok = extract(b'{"a":1,"b":[2,3]}', filename="c.json")
    assert '"a": 1' in ok.text

    broken = extract(b'{"a":1,', filename="c.json")
    assert "not valid JSON" in " ".join(broken.notes)
    assert '{"a":1,' in broken.text


def test_html_drops_script_bodies_not_just_tags() -> None:
    payload = b"<html><body><script>var x=1;</script><p>Real copy</p></body></html>"
    result = extract(payload, filename="page.html")
    assert "Real copy" in result.text
    assert "var x" not in result.text


def test_non_utf8_is_decoded_and_the_guess_is_disclosed() -> None:
    payload = "name,city\nJosé,Málaga\n".encode("latin-1")
    result = extract(payload, filename="people.csv")
    assert "José" in result.text or "Jos" in result.text
    assert any("UTF-8" in n for n in result.notes)


def test_legacy_formats_are_refused_by_name() -> None:
    """A refusal has to tell the sender what to do about it."""
    with pytest.raises(UnsupportedFileTypeError) as excinfo:
        extract(b"\xd0\xcf\x11\xe0", filename="budget.xls")
    assert ".xlsx" in excinfo.value.reason


def test_unknown_format_names_what_it_saw() -> None:
    with pytest.raises(UnsupportedFileTypeError) as excinfo:
        extract(b"\x00\x01", filename="thing.qqq")
    assert "thing.qqq" in excinfo.value.reason


def test_oversize_payload_is_refused_before_parsing() -> None:
    with pytest.raises(FileTooLargeError) as excinfo:
        extract(b"x" * (26 * 1024 * 1024), filename="huge.csv")
    assert "not read" in excinfo.value.reason


def test_image_is_acknowledged_but_never_claimed_as_read() -> None:
    result = extract(b"\x89PNG\r\n", filename="chart.png")
    assert result.kind == "image"
    assert "NOT looked at" in result.text
    assert "Do NOT guess" in result.text


def test_xlsx_reads_every_sheet() -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    first = workbook.active
    assert first is not None
    first.title = "Pipeline"
    first.append(["District", "Stage"])
    first.append(["Austin ISD", "Demo"])
    second = workbook.create_sheet("Notes")
    second.append(["Owner", "Comment"])
    second.append(["Josh", "follow up"])

    buffer = io.BytesIO()
    workbook.save(buffer)

    result = extract(buffer.getvalue(), filename="pipeline.xlsx")
    assert [t.sheet_name for t in result.tables] == ["Pipeline", "Notes"]
    assert result.tables[0].columns == ["District", "Stage"]
    assert "Notes" in " ".join(result.notes)


def test_docx_includes_tables_not_only_paragraphs() -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("Q3 plan")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Channel"
    table.cell(0, 1).text = "Spend"
    table.cell(1, 0).text = "Paid search"
    table.cell(1, 1).text = "42000"

    buffer = io.BytesIO()
    document.save(buffer)

    result = extract(buffer.getvalue(), filename="plan.docx")
    assert "Q3 plan" in result.text
    assert "Paid search" in result.text
    assert "42000" in result.text


def test_supported_extensions_covers_the_common_marketing_formats() -> None:
    assert {".csv", ".tsv", ".xlsx", ".pdf", ".docx"} <= supported_extensions()
