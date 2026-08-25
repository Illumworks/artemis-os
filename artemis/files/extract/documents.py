"""PDF and Word extraction.

PDF reuses `artemis.scouts._pdf.extract_text`, which already backs the board-
minutes and state-DoE scouts -- including its OCR fallback for scanned pages.
Reusing it means a scanned PDF dropped into Slack behaves exactly like one
pulled off a district website, and there is only one PDF path to maintain.
"""

from __future__ import annotations

import io

from artemis.files.extract.base import ExtractedFile, FileParseError, TabularShape, cap_text


def extract_pdf(
    payload: bytes, *, filename: str, mimetype: str = "application/pdf"
) -> ExtractedFile:
    """Extract a PDF's text, falling back to OCR for scanned pages."""
    from artemis.scouts._pdf import extract_text as pdf_text

    try:
        raw = pdf_text(payload)
    except Exception as exc:
        raise FileParseError(
            f"{filename} could not be read as a PDF ({type(exc).__name__}). "
            "If it is password-protected it must be unlocked first."
        ) from exc

    notes: list[str] = []
    if not raw.strip():
        # An empty result is genuinely ambiguous and must not read as "the file
        # is blank". A scanned PDF with no OCR available produces exactly this.
        raise FileParseError(
            f"{filename} opened as a PDF but no text could be extracted. It is most "
            "likely a scan or images-only; OCR either is not installed or could not "
            "read it. The file is not necessarily empty."
        )

    text, truncated = cap_text(raw)
    return ExtractedFile(
        filename=filename,
        mimetype=mimetype,
        kind="document",
        text=text,
        size_bytes=len(payload),
        truncated=truncated,
        notes=notes,
    )


def extract_docx(payload: bytes, *, filename: str, mimetype: str = "") -> ExtractedFile:
    """Extract a Word document's paragraphs AND its tables.

    Tables are pulled out separately because Word documents routinely carry the
    real payload in a table -- a schedule, a budget, a matrix -- and a paragraph-
    only walk silently drops every one of them, returning a confident, fluent,
    incomplete document.
    """
    try:
        import docx

        document = docx.Document(io.BytesIO(payload))
    except Exception as exc:
        raise FileParseError(
            f"{filename} could not be opened as a Word document ({type(exc).__name__}). "
            "Note the legacy .doc format is not supported -- only .docx."
        ) from exc

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    shapes: list[TabularShape] = []
    rendered_tables: list[str] = []
    for index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rows = [r for r in rows if any(r)]
        if not rows:
            continue
        header, *body = rows
        shapes.append(
            TabularShape(
                columns=header,
                total_rows=len(body),
                sample_rows=body,
                sheet_name=f"table {index}",
            )
        )
        rendered_tables.append(f"[table {index}]\n" + "\n".join(" | ".join(row) for row in rows))

    if not paragraphs and not rendered_tables:
        raise FileParseError(f"{filename} opened as a Word document but contains no text.")

    body_text = "\n\n".join(["\n".join(paragraphs), *rendered_tables]).strip()
    text, truncated = cap_text(body_text)

    notes: list[str] = []
    if shapes:
        notes.append(f"Contains {len(shapes)} table(s), included below the prose.")

    return ExtractedFile(
        filename=filename,
        mimetype=mimetype
        or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        kind="document",
        text=text,
        size_bytes=len(payload),
        tables=shapes,
        truncated=truncated,
        notes=notes,
    )


def describe_image(payload: bytes, *, filename: str, mimetype: str = "") -> ExtractedFile:
    """Acknowledge an image without pretending to have read it.

    Vision is a deliberate later slice. Until then this returns a factual
    placeholder rather than raising, so an agent can say "you sent me a
    screenshot, I can see it is there but cannot read it yet" -- which is true,
    useful, and cannot be mistaken for having understood the contents.
    """
    return ExtractedFile(
        filename=filename,
        mimetype=mimetype or "image/unknown",
        kind="image",
        text=(
            f"[image: {filename}, {len(payload):,} bytes] "
            "Image contents cannot be read yet -- this agent has no vision capability "
            "wired up. Acknowledge the image and ask the sender to describe it or paste "
            "the relevant text. Do NOT guess at what it shows."
        ),
        size_bytes=len(payload),
        notes=["Image understanding is not enabled; only the file's presence is known."],
    )
